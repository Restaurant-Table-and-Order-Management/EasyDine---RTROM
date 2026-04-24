from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_manual():
    doc = Document()

    # Title Page
    title = doc.add_heading('EasyDine RTROM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('User Manual & Technical Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        "1. Quick Project Summary",
        "2. Versions We Are Using",
        "3. Tools — What They Are & Why We Need Them",
        "4. Installation — Step by Step with Verification",
        "5. Clone the Repositories",
        "6. Running the Backend (Spring Boot)",
        "7. Backend Troubleshooting",
        "8. Running the Frontend (React/Vite)",
        "9. Connect Frontend to Backend",
        "10. Troubleshooting & Rescue Plan",
        "11. Verification Checklist",
        "12. Post-Startup Initialization",
        "13. Application Configuration",
        "14. Updating Default Admin User Details"
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_page_break()

    # 1. Quick Project Summary
    doc.add_heading('1. Quick Project Summary', level=1)
    doc.add_paragraph(
        "EasyDine RTROM (Restaurant Table Reservation & Operations Management) is a premium, "
        "production-grade SaaS solution designed to bridge the gap between traditional dining "
        "operations and modern technology. The platform provides a seamless ecosystem for "
        "restaurant owners, kitchen staff, and customers, focusing on real-time synchronization, "
        "data-driven financial insights, and high-end user experiences."
    )
    doc.add_paragraph("Key Feature Pillars:")
    doc.add_paragraph("• Real-time Kitchen-Waiter Synchronization", style='List Bullet')
    doc.add_paragraph("• Seamless Digital Reservation System", style='List Bullet')
    doc.add_paragraph("• Integrated Finance Ledger & Analytics Dashboard", style='List Bullet')
    doc.add_paragraph("• Role-Based Access Control (Admin, Kitchen, Staff, Customer)", style='List Bullet')

    # 2. Versions We Are Using
    doc.add_heading('2. Versions We Are Using', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Version'
    
    rows = [
        ('Java', '21 (LTS)'),
        ('Spring Boot', '3.2.4'),
        ('React', '18.2.0'),
        ('Vite', '5.0.2'),
        ('Node.js', '18.0.0+'),
        ('MySQL', '8.0.33'),
        ('Maven', '3.8+')
    ]
    for comp, ver in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = ver

    # 3. Tools — What They Are & Why We Need Them
    doc.add_heading('3. Tools — What They Are & Why We Need Them', level=1)
    tools = [
        ("JDK 21", "The Java Development Kit is the foundation of the backend. It provides the compiler and runtime for Spring Boot."),
        ("Node.js & NPM", "Essential for the frontend development environment. It manages packages and runs the Vite build tool."),
        ("MySQL Server", "The relational database used for persistent storage of users, reservations, and financial data."),
        ("Maven", "Handles project build lifecycle, dependency management, and automated testing for the backend."),
        ("VS Code / IntelliJ", "Recommended IDEs for full-stack development with React and Java support.")
    ]
    for tool, desc in tools:
        p = doc.add_paragraph()
        p.add_run(f"{tool}: ").bold = True
        p.add_run(desc)

    # 4. Installation — Step by Step with Verification
    doc.add_heading('4. Installation — Step by Step with Verification', level=1)
    doc.add_paragraph("Follow these steps in order to ensure a stable local environment:")
    doc.add_paragraph("1. Verify Prerequisites: Run 'java -version' and 'node -v' in your terminal.", style='List Number')
    doc.add_paragraph("2. Database Creation: Log in to MySQL and execute: 'CREATE DATABASE easydine;'", style='List Number')
    doc.add_paragraph("3. Environment Config: Duplicate the '.env.example' files in both backend and frontend roots.", style='List Number')

    # 5. Clone the Repositories
    doc.add_heading('5. Clone the Repositories', level=1)
    doc.add_paragraph("Use Git to pull the latest production source code:")
    doc.add_paragraph("git clone https://github.com/your-repo/easydine-rtrom.git", style='No Spacing')
    doc.add_paragraph("Ensure you have permissions to access the repository and that your SSH keys are configured.")

    # 6. Running the Backend (Spring Boot)
    doc.add_heading('6. Running the Backend (Spring Boot)', level=1)
    doc.add_paragraph("1. Navigate to the backend directory: 'cd backend'", style='List Number')
    doc.add_paragraph("2. Install dependencies: 'mvn clean install'", style='List Number')
    doc.add_paragraph("3. Launch application: 'mvn spring-boot:run'", style='List Number')
    doc.add_paragraph("The server should start on port 8080 by default.")

    # 7. Backend Troubleshooting
    doc.add_heading('7. Backend Troubleshooting', level=1)
    doc.add_paragraph("Common Issues:")
    doc.add_paragraph("• DB Connection Refused: Verify MySQL is running and your .env credentials match.", style='List Bullet')
    doc.add_paragraph("• Port 8080 Busy: Kill existing processes using 'lsof -i :8080' or change port in application.properties.", style='List Bullet')
    doc.add_paragraph("• Maven Plugins Fail: Ensure your Maven version is compatible with Java 21.", style='List Bullet')

    # 8. Running the Frontend (React/Vite)
    doc.add_heading('8. Running the Frontend (React/Vite)', level=1)
    doc.add_paragraph("1. Navigate to the frontend directory: 'cd frontend'", style='List Number')
    doc.add_paragraph("2. Install packages: 'npm install'", style='List Number')
    doc.add_paragraph("3. Start Dev Server: 'npm run dev'", style='List Number')
    doc.add_paragraph("The application will be accessible at http://localhost:5173.")

    # 9. Connect Frontend to Backend
    doc.add_heading('9. Connect Frontend to Backend', level=1)
    doc.add_paragraph(
        "For the frontend to communicate with the backend, verify the API configuration in "
        "'frontend/src/api/axios.config.js' (or your respective auth store). The base URL should "
        "point to 'http://localhost:8080/api/v1'."
    )

    # 10. Troubleshooting & Rescue Plan
    doc.add_heading('10. Troubleshooting & Rescue Plan', level=1)
    doc.add_paragraph("• CORS Policy Errors: Ensure '@CrossOrigin' is configured correctly in Spring Boot Controllers.", style='List Bullet')
    doc.add_paragraph("• Missing UI Elements: Check browser console (F12) for failed component imports.", style='List Bullet')
    doc.add_paragraph("• Node Modules Issues: Delete 'node_modules' folder and 'package-lock.json', then run 'npm install'.", style='List Bullet')

    # 11. Verification Checklist
    doc.add_heading('11. Verification Checklist', level=1)
    doc.add_paragraph("[ ] Backend starts without stack traces", style='List Bullet')
    doc.add_paragraph("[ ] Database tables are auto-generated on first run", style='List Bullet')
    doc.add_paragraph("[ ] Landing page loads correctly with all assets", style='List Bullet')
    doc.add_paragraph("[ ] User login results in a successful JWT token storage", style='List Bullet')

    # 12. Post-Startup Initialization
    doc.add_heading('12. Post-Startup Initialization', level=1)
    doc.add_paragraph(
        "Upon first login, the system will check for a default admin. If not present, "
        "the 'DataInitializer' service in the backend will provision the initial "
        "ADMIN and KITCHEN roles. Ensure you update these immediately."
    )

    # 13. Application Configuration
    doc.add_heading('13. Application Configuration', level=1)
    doc.add_paragraph("Configuration Options:")
    doc.add_paragraph("• Theme Settings: Accessible via the Navbar toggle (Dark/Light Mode).", style='List Bullet')
    doc.add_paragraph("• System Roles: Defined in the 'Role' Entity (ADMIN, KITCHEN, STAFF, CUSTOMER).", style='List Bullet')

    # 14. Updating Default Admin User Details
    doc.add_heading('14. Updating Default Admin User Details', level=1)
    doc.add_paragraph(
        "1. Log in with the default admin credentials (e.g., admin@easydine.com / admin123).\n"
        "2. Navigate to 'Settings' or 'Profile' from the sidebar.\n"
        "3. Enter your new Email and Password.\n"
        "4. Click 'Save Changes' to update the secure hash in the database."
    )

    # Save the document
    doc.save('EasyDine_User_Manual.docx')
    print("Manual generated successfully as EasyDine_User_Manual.docx")

if __name__ == "__main__":
    create_manual()
